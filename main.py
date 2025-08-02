import pytesseract
from pdf2image import convert_from_path
from docx import Document
from PIL import Image
import os
import sys

# Optional: set path to tesseract if not in system PATH
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def image_to_text(image_path):
    image = Image.open(image_path)
    return pytesseract.image_to_string(image)

def pdf_to_text(pdf_path):
    pages = convert_from_path(pdf_path)
    texts = []
    for i, page in enumerate(pages):
        print(f"[INFO] Processing page {i+1}")
        texts.append(pytesseract.image_to_string(page))
    return texts

def save_to_word(text_list, output_path="output.docx"):
    doc = Document()
    for i, text in enumerate(text_list):
        doc.add_heading(f'Page {i+1}', level=2)
        doc.add_paragraph(text)
    doc.save(output_path)
    print(f"[SUCCESS] Saved output to: {output_path}")

def main():
    if len(sys.argv) < 2:
        print("Usage: python main.py <file>")
        return

    file_path = sys.argv[1]

    if not os.path.exists(file_path):
        print("File not found!")
        return

    if file_path.lower().endswith(".pdf"):
        text = pdf_to_text(file_path)
    else:
        text = [image_to_text(file_path)]

    save_to_word(text)

if __name__ == "__main__":
    main()
